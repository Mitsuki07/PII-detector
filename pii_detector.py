import re
import os
import fitz #pymupdf
import pytesseract #ocr
import pandas as pd
from PIL import Image
from docx import Document

def text_from_pdf(pdf_path):        #function to extract text from PDF
    try:                            
        text = ''
        with fitz.open(pdf_path) as doc:
             for page in doc:
                text += page.get_text("text")
        return text
    except Exception as e:
        print(f"Error processing PDF: {e}")
        return ""

def text_from_image(image_path):    #function to extract text from img
    pytesseract.pytesseract.tesseract_cmd = 'C:/Program Files/Tesseract-OCR/tesseract.exe'
    try:
        image = Image.open(image_path)
        return pytesseract.image_to_string(image)
    except Exception as e:
        print(f"Error processing image: {e}")
        return ""
    
def text_from_docx(docx_path):      #function to extract txt from docs
    try:
        text = ""
        doc = Document(docx_path)
        
        #extract text from paragraphs
        for para in doc.paragraphs:
            text += para.text + "\n"

        #extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
            text += "\n"
        return text.strip()
    except Exception as e:
        print(f"Error processing docx: {e}")
        return ""

def text_from_excel(excel_path):    #func to extract txt from excel
    try:
        ext = os.path.splitext(excel_path)[1].lower()
        if ext == ".csv":
            df = pd.read_csv(excel_path)
        else:
            df = pd.read_excel(excel_path, engine="openpyxl")

        return "\n".join(df.astype(str).stack().tolist())
    except Exception as e:
        print(f"Error processing excel: {e}")
        return ""
    
def extract_pii(text):
    try:
        patterns = {
            "Email": r'[\w\.-]+@[\w\.-]+\.\w+',
            "Phone": r'\b(?:\+91[-\s]?)?[6789]\d{9}\b',
            "Aadhar": r'\b\d{4}\s\d{4}\s\d{4}\b',
            "PAN": r'[A-Z]{5}[0-9]{4}[A-Z]{1}'
        }
        redacted_pii = {}

        for key, pattern in patterns.items():
            matches = re.findall(pattern, text)
            redacted_values = []
            for match in matches:
                if key == "Email":
                    redacted_values.append("*****@example.com")
                elif key == "Phone":
                    redacted_values.append("**********")
                elif key == "Aadhar":
                    redacted_values.append("**** **** ****")
                elif key == "PAN":
                    redacted_values.append("*****1234*")
            redacted_pii[key] = redacted_values
        return redacted_pii
    
    except Exception as e:
        print(f"Error extracting PII: {e}")
        return {}, text        

def process_file(file_path):
    try:
        ext = os.path.splitext(file_path)[1].lower()
        if ext == ".pdf":
            text = text_from_pdf(file_path)
        elif ext in [".jpg", ".jpeg", ".png"]:
            text = text_from_image(file_path)
        elif ext == ".docx":
            text = text_from_docx(file_path)
        elif ext in [".xls", ".xlsx", ".csv"]:
            text = text_from_excel(file_path)
        elif ext == ".txt":
            with open(file_path, "r") as f:
                text = f.read()
        else:
            print("Unsupported file type")
            return {}, ""
        
        return extract_pii(text)
    except Exception as e:
        print(f"Error processing file: {e}")
        return {}, ""
    
def process_directory(directory_path):
    pii_results_all = {}
    for root, _, files in os.walk(directory_path):
        for file in files:
            file_path = os.path.join(root, file)
            redacted_pii = process_file(file_path)
            pii_results_all[file_path] = redacted_pii
            print(f"Processed: {file_path}")
            
            # Check if any PII was detected
            if any(redacted_pii.values()):  # If PII is found
                print(f"Redacted PII: {redacted_pii}")
            else:  # If no PII is found
                print("No PII found.")
    return pii_results_all
    
if __name__ == "__main__":
    directory_path = "./sample_files" 
    pii_results_all = process_directory(directory_path)
    
